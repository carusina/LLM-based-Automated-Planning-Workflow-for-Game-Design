# models/document_generator.py
import os
import json
import datetime
from typing import Dict, Any, List, Optional, Union
import markdown
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import docx

class DocumentGenerator:
    """게임 기획서 문서 생성기"""
    
    def __init__(self, output_dir: str = "output"):
        """
        문서 생성기 초기화
        
        Args:
            output_dir: 문서 출력 디렉토리
        """
        self.output_dir = output_dir
        
        # 출력 디렉토리가 없으면 생성
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
    
    def create_document(self, 
                       game_design: Dict[str, Any], 
                       format: str = "markdown",
                       filename: str = None) -> str:
        """
        게임 기획서 문서 생성
        
        Args:
            game_design: 게임 기획 정보
            format: 출력 형식 ("markdown", "pdf", "docx")
            filename: 파일 이름 (지정하지 않으면 자동 생성)
            
        Returns:
            생성된 문서 경로
        """
        # 파일 이름이 지정되지 않은 경우 자동 생성
        if not filename:
            game_title = game_design.get("game_title", "게임기획서")
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{game_title.replace(' ', '_')}_{timestamp}"
        
        # 파일 경로 생성
        print(f"생성 요청된 형식: {format}")
        print(f"생성될 파일 경로: {os.path.join(self.output_dir, f'{filename}.*')}")
        
        try:
            if format.lower() == "markdown":
                file_path = os.path.join(self.output_dir, f"{filename}.md")
                self._create_markdown(game_design, file_path)
                print(f"마크다운 파일 생성 완료: {file_path}")
            elif format.lower() == "pdf":
                file_path = os.path.join(self.output_dir, f"{filename}.pdf")
                self._create_pdf(game_design, file_path)
                print(f"PDF 파일 생성 완료: {file_path}")
            elif format.lower() == "docx":
                file_path = os.path.join(self.output_dir, f"{filename}.docx")
                self._create_docx(game_design, file_path)
                print(f"Word 문서 생성 완료: {file_path}")
            else:
                raise ValueError(f"지원되지 않는 형식: {format}")
        except Exception as e:
            print(f"문서 생성 중 오류 발생: {e}")
            raise
        
        return file_path
    
    def _create_markdown(self, game_design: Dict[str, Any], file_path: str) -> None:
        """
        마크다운 형식으로 게임 기획서 생성
        
        Args:
            game_design: 게임 기획 정보
            file_path: 출력 파일 경로
        """
        with open(file_path, 'w', encoding='utf-8') as f:
            # 제목 및 기본 정보
            f.write(f"# {game_design.get('game_title', '게임 기획서')}\n\n")
            f.write(f"**고수준 컨셉:** {game_design.get('high_concept', '')}\n\n")
            f.write(f"**타겟 사용자층:** {game_design.get('target_audience', '')}\n")
            f.write(f"**장르:** {game_design.get('genre', '')}\n")
            f.write(f"**플랫폼:** {game_design.get('platform', '')}\n\n")
            
            # 컨셉 세부사항
            f.write("## 게임 컨셉\n\n")
            concept_details = game_design.get("concept_details", {})
            f.write(f"### 확장된 컨셉\n{concept_details.get('extended_concept', '')}\n\n")
            
            f.write("### 차별화 포인트\n")
            for usp in concept_details.get("unique_selling_points", []):
                f.write(f"- {usp}\n")
            f.write("\n")
            
            f.write(f"### 분위기\n{concept_details.get('mood', '')}\n\n")
            
            # 게임플레이
            f.write("## 게임플레이\n\n")
            gameplay = game_design.get("gameplay", {})
            f.write(f"### 핵심 게임플레이 루프\n{gameplay.get('core_gameplay_loop', '')}\n\n")
            
            f.write("### 플레이어 액션\n")
            for action in gameplay.get("player_actions", []):
                f.write(f"- {action}\n")
            f.write("\n")
            
            f.write(f"### 진행 시스템\n{gameplay.get('progression_system', '')}\n\n")
            
            f.write("### 도전 유형\n")
            for challenge in gameplay.get("challenge_types", []):
                f.write(f"- {challenge}\n")
            f.write("\n")
            
            f.write("### 보상 시스템\n")
            for reward in gameplay.get("reward_systems", []):
                f.write(f"- {reward}\n")
            f.write("\n")
            
            f.write("### 게임 모드\n")
            for mode in gameplay.get("game_modes", []):
                f.write(f"- {mode}\n")
            f.write("\n")
            
            f.write(f"### 조작 방식\n{gameplay.get('controls', '')}\n\n")
            
            f.write("### 독특한 메커니즘\n")
            for mechanic in gameplay.get("unique_mechanics", []):
                f.write(f"- {mechanic}\n")
            f.write("\n")
            
            # 내러티브
            f.write("## 내러티브\n\n")
            narrative = game_design.get("narrative", {})
            f.write(f"### 세계관\n{narrative.get('setting', '')}\n\n")
            f.write(f"### 배경 스토리\n{narrative.get('background_lore', '')}\n\n")
            f.write(f"### 주요 스토리라인\n{narrative.get('main_plot', '')}\n\n")
            f.write(f"### 스토리 구조\n{narrative.get('plot_structure', '')}\n\n")
            
            f.write("### 주요 테마\n")
            for theme in narrative.get("themes", []):
                f.write(f"- {theme}\n")
            f.write("\n")
            
            f.write("### 캐릭터\n")
            for character in narrative.get("characters", []):
                f.write(f"#### {character.get('name', '이름 없음')}\n")
                f.write(f"- **역할:** {character.get('role', '')}\n")
                f.write(f"- **배경:** {character.get('background', '')}\n")
                f.write(f"- **동기:** {character.get('motivation', '')}\n")
                f.write(f"- **캐릭터 아크:** {character.get('arc', '')}\n\n")
            
            f.write("### 내러티브 장치\n")
            for device in narrative.get("narrative_devices", []):
                f.write(f"- {device}\n")
            f.write("\n")
            
            f.write(f"### 플레이어 선택\n{narrative.get('player_agency', '')}\n\n")
            f.write(f"### 게임플레이와 내러티브 통합\n{narrative.get('integration_with_gameplay', '')}\n\n")
            
            # 아트 디렉션
            f.write("## 아트 디렉션\n\n")
            art = game_design.get("art_direction", {})
            f.write(f"### 시각적 스타일\n{art.get('visual_style', '')}\n\n")
            f.write(f"### 색상 팔레트\n{art.get('color_palette', '')}\n\n")
            
            f.write("### 참고 아트 스타일\n")
            for ref in art.get("art_references", []):
                f.write(f"- {ref}\n")
            f.write("\n")
            
            f.write(f"### 캐릭터 디자인\n{art.get('character_design', '')}\n\n")
            f.write(f"### 환경 디자인\n{art.get('environment_design', '')}\n\n")
            f.write(f"### UI 디자인\n{art.get('ui_design', '')}\n\n")
            f.write(f"### 애니메이션 스타일\n{art.get('animation_style', '')}\n\n")
            f.write(f"### 조명\n{art.get('lighting', '')}\n\n")
            f.write(f"### 사운드 디자인\n{art.get('sound_design', '')}\n\n")
            f.write(f"### 음악 방향\n{art.get('music_direction', '')}\n\n")
            
            # 기술 사양
            f.write("## 기술 사양\n\n")
            tech = game_design.get("technical_specs", {})
            
            f.write("### 타겟 플랫폼\n")
            for platform in tech.get("target_platforms", []):
                f.write(f"- {platform}\n")
            f.write("\n")
            
            f.write(f"### 최소 사양\n{tech.get('minimum_specs', '')}\n\n")
            f.write(f"### 권장 사양\n{tech.get('recommended_specs', '')}\n\n")
            f.write(f"### 게임 엔진\n{tech.get('engine', '')}\n\n")
            
            f.write("### 핵심 기술\n")
            for technology in tech.get("key_technologies", []):
                f.write(f"- {technology}\n")
            f.write("\n")
            
            f.write(f"### 네트워킹\n{tech.get('networking', '')}\n\n")
            f.write(f"### 성능 목표\n{tech.get('performance_targets', '')}\n\n")
            
            f.write("### 개발 도전 과제\n")
            for challenge in tech.get("development_challenges", []):
                f.write(f"- {challenge}\n")
            f.write("\n")
            
            f.write(f"### 에셋 요구사항\n{tech.get('asset_requirements', '')}\n\n")
            f.write(f"### 예상 개발 리소스\n{tech.get('estimated_development_resources', '')}\n\n")
            
            # 수익화 계획
            f.write("## 수익화 계획\n\n")
            monetization = game_design.get("monetization", {})
            f.write(f"### 수익화 모델\n{monetization.get('monetization_model', '')}\n\n")
            f.write(f"### 가격 책정\n{monetization.get('price_point', '')}\n\n")
            f.write(f"### 인앱 구매\n{monetization.get('in_app_purchases', '')}\n\n")
            f.write(f"### 가상 화폐\n{monetization.get('virtual_currency', '')}\n\n")
            f.write(f"### 배틀패스\n{monetization.get('battle_pass', '')}\n\n")
            f.write(f"### 광고\n{monetization.get('advertising', '')}\n\n")
            f.write(f"### DLC/확장팩\n{monetization.get('dlc_expansion', '')}\n\n")
            f.write(f"### 구독\n{monetization.get('subscription', '')}\n\n")
            
            f.write("### 플레이어 유지 전략\n")
            for strategy in monetization.get("player_retention_strategies", []):
                f.write(f"- {strategy}\n")
            f.write("\n")
            
            f.write(f"### 수익 예상\n{monetization.get('revenue_projections', '')}\n\n")
            f.write(f"### 시장 분석\n{monetization.get('market_analysis', '')}\n\n")
            
            # 개발 로드맵
            f.write("## 개발 로드맵\n\n")
            roadmap = game_design.get("development_roadmap", {})
            timeline = roadmap.get("development_timeline", {})
            
            f.write("### 개발 타임라인\n")
            f.write(f"- **사전 제작:** {timeline.get('pre_production', '')}\n")
            f.write(f"- **프로토타입:** {timeline.get('prototype', '')}\n")
            f.write(f"- **본 제작:** {timeline.get('production', '')}\n")
            f.write(f"- **알파:** {timeline.get('alpha', '')}\n")
            f.write(f"- **베타:** {timeline.get('beta', '')}\n")
            f.write(f"- **골드 마스터:** {timeline.get('gold', '')}\n")
            f.write(f"- **출시 후 지원:** {timeline.get('post_launch', '')}\n\n")
            
            team = roadmap.get("team_structure", {})
            f.write("### 팀 구성\n")
            
            f.write("#### 핵심 팀\n")
            for role in team.get("core_team", []):
                f.write(f"- {role}\n")
            f.write("\n")
            
            f.write("#### 확장 팀\n")
            for role in team.get("expanded_team", []):
                f.write(f"- {role}\n")
            f.write("\n")
            
            f.write("#### 외주\n")
            for outsource in team.get("outsourcing", []):
                f.write(f"- {outsource}\n")
            f.write("\n")
            
            f.write(f"### 예산 추정\n{roadmap.get('budget_estimate', '')}\n\n")
            
            f.write("### 리스크 평가\n")
            for risk in roadmap.get("risk_assessment", []):
                f.write(f"- {risk}\n")
            f.write("\n")
            
            f.write(f"### 테스트 계획\n{roadmap.get('testing_plan', '')}\n\n")
            f.write(f"### 마케팅 타임라인\n{roadmap.get('marketing_timeline', '')}\n\n")
            
            f.write("### 주요 산출물\n")
            for deliverable in roadmap.get("key_deliverables", []):
                f.write(f"- {deliverable}\n")
            f.write("\n")
            
            f.write("### 성공 지표\n")
            for metric in roadmap.get("success_metrics", []):
                f.write(f"- {metric}\n")
            f.write("\n")
            
            # 경쟁 분석 (있는 경우)
            if "competition_analysis" in game_design:
                f.write("## 경쟁 분석\n\n")
                competition = game_design.get("competition_analysis", {})
                
                for i, competitor in enumerate(competition.get("competitors", []), 1):
                    f.write(f"### 경쟁작 {i}: {competitor.get('game_name', '이름 없음')}\n\n")
                    f.write(f"**개발사:** {competitor.get('developer', '')}\n")
                    f.write(f"**퍼블리셔:** {competitor.get('publisher', '')}\n")
                    f.write(f"**출시일:** {competitor.get('release_date', '')}\n")
                    f.write(f"**장르:** {competitor.get('genre', '')}\n")
                    f.write(f"**타겟 사용자층:** {competitor.get('target_audience', '')}\n\n")
                    
                    f.write("#### 주요 특징\n")
                    for feature in competitor.get("key_features", []):
                        f.write(f"- {feature}\n")
                    f.write("\n")
                    
                    f.write(f"#### 게임플레이 분석\n{competitor.get('gameplay_analysis', '')}\n\n")
                    f.write(f"#### 내러티브 분석\n{competitor.get('narrative_analysis', '')}\n\n")
                    f.write(f"#### 아트 스타일 분석\n{competitor.get('art_style_analysis', '')}\n\n")
                    f.write(f"#### 수익화 모델\n{competitor.get('monetization_model', '')}\n\n")
                    
                    reception = competitor.get("reception", {})
                    f.write("#### 시장 반응\n")
                    f.write(f"- **평론가 점수:** {reception.get('critic_score', '')}\n")
                    f.write(f"- **사용자 평점:** {reception.get('user_score', '')}\n")
                    f.write(f"- **판매 실적:** {reception.get('sales_performance', '')}\n\n")
                    
                    f.write("#### 강점\n")
                    for strength in competitor.get("strengths", []):
                        f.write(f"- {strength}\n")
                    f.write("\n")
                    
                    f.write("#### 약점\n")
                    for weakness in competitor.get("weaknesses", []):
                        f.write(f"- {weakness}\n")
                    f.write("\n")
                    
                    f.write("#### 배울 점\n")
                    for lesson in competitor.get("lessons", []):
                        f.write(f"- {lesson}\n")
                    f.write("\n")

            # 스토리라인 정보가 있는 경우
            if "storyline" in game_design:
                storyline = game_design.get("storyline", {})
                
                f.write("## 스토리라인\n\n")
                f.write(f"### 스토리 제목\n{storyline.get('title', '')}\n\n")
                f.write(f"### 전제\n{storyline.get('premise', '')}\n\n")
                f.write(f"### 스토리 아크\n{storyline.get('story_arc', '')}\n\n")
                
                f.write("### 주요 테마\n")
                for theme in storyline.get('themes', []):
                    f.write(f"- {theme}\n")
                f.write("\n")
                
                f.write("### 주요 플롯 포인트\n")
                for plot_point in storyline.get('major_plot_points', []):
                    f.write(f"- {plot_point}\n")
                f.write("\n")
                
                f.write("### 가능한 엔딩\n")
                for ending in storyline.get('possible_endings', []):
                    f.write(f"- {ending}\n")
                f.write("\n")
                
                # 챕터 개요
                f.write("### 챕터 개요\n\n")
                for chapter in storyline.get('chapter_outlines', []):
                    chapter_num = chapter.get('chapter_number', '')
                    chapter_title = chapter.get('title', '')
                    chapter_synopsis = chapter.get('synopsis', '')
                    
                    f.write(f"#### 챕터 {chapter_num}: {chapter_title}\n")
                    f.write(f"{chapter_synopsis}\n\n")
                    
                    f.write("**목표:**\n")
                    for goal in chapter.get('goals', []):
                        f.write(f"- {goal}\n")
                    f.write("\n")
                    
                    f.write("**주요 위치:**\n")
                    for location in chapter.get('key_locations', []):
                        f.write(f"- {location}\n")
                    f.write("\n")
                    
                    f.write("**주요 사건:**\n")
                    for event in chapter.get('key_events', []):
                        f.write(f"- {event}\n")
                    f.write("\n")
                    
                    f.write("**도전 과제:**\n")
                    for challenge in chapter.get('challenges', []):
                        f.write(f"- {challenge}\n")
                    f.write("\n")
                
                # 챕터 상세 내용
                f.write("### 챕터 상세 내용\n\n")
                for chapter in storyline.get('chapter_details', []):
                    chapter_num = chapter.get('chapter_number', '')
                    chapter_title = chapter.get('title', '')
                    
                    f.write(f"#### 챕터 {chapter_num}: {chapter_title}\n")
                    f.write(f"{chapter.get('detailed_synopsis', '')}\n\n")
                    
                    f.write("**오프닝 씬:**\n")
                    f.write(f"{chapter.get('opening_scene', '')}\n\n")
                    
                    f.write("**주요 사건:**\n")
                    for event in chapter.get('key_events', []):
                        f.write(f"- **{event.get('event_title', '')}**: {event.get('description', '')}\n")
                    f.write("\n")
                    
                    f.write("**클라이맥스:**\n")
                    f.write(f"{chapter.get('climax', '')}\n\n")
                    
                    f.write("**엔딩:**\n")
                    f.write(f"{chapter.get('ending', '')}\n\n")
                
                # 스토리 분기
                f.write("### 스토리 분기\n\n")
                for branch in storyline.get('story_branches', {}).get('story_branches', []):
                    f.write(f"#### {branch.get('branch_name', '')}\n")
                    f.write(f"**트리거:** {branch.get('trigger', '')}\n")
                    f.write(f"**시놉시스:** {branch.get('synopsis', '')}\n\n")
                    
                    f.write("**주요 사건:**\n")
                    for event in branch.get('key_events', []):
                        f.write(f"- {event}\n")
                    f.write("\n")
                    
                    f.write("**결말:**\n")
                    f.write(f"{branch.get('ending', '')}\n\n")
                
                # 샘플 대화 스크립트
                if 'sample_dialogue' in storyline:
                    dialogue = storyline.get('sample_dialogue', {})
                    
                    f.write("### 샘플 대화 스크립트\n\n")
                    f.write(f"**씬:** {dialogue.get('scene_title', '')}\n")
                    f.write(f"**배경:** {dialogue.get('scene_setting', '')}\n")
                    f.write(f"**톤:** {dialogue.get('tone', '')}\n\n")
                    
                    f.write("**스크립트:**\n")
                    for line in dialogue.get('script', []):
                        speaker = line.get('speaker', '')
                        text = line.get('line', '')
                        emotion = line.get('emotion', '')
                        action = line.get('action', '')
                        
                        if action:
                            f.write(f"*{speaker} ({emotion})* [{action}]: {text}\n")
                        else:
                            f.write(f"*{speaker} ({emotion})*: {text}\n")
                    f.write("\n")
                    
                    f.write("**플레이어 선택지:**\n")
                    for option in dialogue.get('player_dialogue_options', []):
                        f.write(f"- **선택:** {option.get('option_text', '')}\n")
                        f.write(f"  **응답:** {option.get('response', '')}\n")
                        f.write(f"  **결과:** {option.get('outcome', '')}\n")
                    f.write("\n")
    
    def _create_pdf(self, game_design: Dict[str, Any], file_path: str) -> None:
        """
        PDF 형식으로 게임 기획서 생성
        
        Args:
            game_design: 게임 기획 정보
            file_path: 출력 파일 경로
        """
        try:
            # 폰트 등록 (한글 폰트 사용 시 필수)
            pdfmetrics.registerFont(TTFont('NanumGothic', '/Users/carusina/programming/SW_Project/assets/fonts/NanumGothic.ttf'))
            pdfmetrics.registerFont(TTFont('NanumGothicBold', '/Users/carusina/programming/SW_Project/assets/fonts/NanumGothicBold.ttf'))
            
            # 마크다운 생성을 통해 내용 확보
            temp_md_path = file_path.replace('.pdf', '_temp.md')
            self._create_markdown(game_design, temp_md_path)
            
            # 마크다운 내용 PDF로 변환
            with open(temp_md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # PDF 생성 라이브러리 사용
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
            from reportlab.lib import colors
            
            # PDF 문서 생성
            pdf = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # 기존 스타일 수정 및 새로운 스타일 생성
            title_style = ParagraphStyle(
                name='Title',
                parent=styles['Title'],
                fontName='NanumGothicBold',  # 볼드 폰트 사용
                fontSize=16
            )
            
            heading2_style = ParagraphStyle(
                name='Heading2',
                parent=styles['Heading2'],
                fontName='NanumGothicBold',
                fontSize=14
            )
            
            heading3_style = ParagraphStyle(
                name='Heading3',
                parent=styles['Heading3'],
                fontName='NanumGothicBold',
                fontSize=12
            )
            
            body_style = ParagraphStyle(
                name='BodyText',
                parent=styles['BodyText'],
                fontName='NanumGothic',
                fontSize=10
            )
            
            # 필요한 추가 스타일만 정의
            bullet_style = ParagraphStyle(
                name='CustomBullet',
                parent=body_style,
                leftIndent=20,
                spaceBefore=3,
                spaceAfter=3,
                fontName='NanumGothic'
            )
            
            bold_style = ParagraphStyle(
                name='CustomBold',
                parent=body_style,
                fontName='NanumGothicBold'
            )
            
            heading4_style = ParagraphStyle(
                name='CustomHeading4',
                parent=heading3_style,
                fontSize=11,
                spaceBefore=6,
                spaceAfter=3,
                fontName='NanumGothicBold'
            )
            
            story = []
            
            # 마크다운 내용을 PDF 문단으로 변환
            lines = md_content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], title_style))
                    story.append(Spacer(1, 12))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], heading2_style))
                    story.append(Spacer(1, 6))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], heading3_style))
                    story.append(Spacer(1, 6))
                elif line.startswith('#### '):
                    story.append(Paragraph(line[5:], heading4_style))
                    story.append(Spacer(1, 3))
                elif line.startswith('- '):
                    story.append(Paragraph('• ' + line[2:], bullet_style))
                elif line.startswith('**') and line.endswith('**'):
                    # 볼드 텍스트 처리
                    text = line.replace('**', '')
                    story.append(Paragraph(text, bold_style))
                elif line:
                    # 일반 텍스트
                    story.append(Paragraph(line, body_style))
                    if i < len(lines) - 1 and lines[i+1].strip():
                        story.append(Spacer(1, 6))
                
                i += 1
                
            # PDF로 저장
            pdf.build(story)
            
            # 임시 마크다운 파일 삭제
            if os.path.exists(temp_md_path):
                os.remove(temp_md_path)
                
            print(f"PDF 파일 생성 완료: {file_path}")
            
        except Exception as e:
            print(f"PDF 생성 중 오류 발생: {str(e)}")
            raise
    
    def _create_docx(self, game_design: Dict[str, Any], file_path: str) -> None:
        """
        Word 문서 형식으로 게임 기획서 생성
        
        Args:
            game_design: 게임 기획 정보
            file_path: 출력 파일 경로
        """
        try:
            # python-docx 라이브러리를 사용하여 Word 문서 생성
            doc = docx.Document()
            
            # 제목 및 기본 정보
            doc.add_heading(game_design.get('game_title', '게임 기획서'), 0)
            
            # 기본 정보 섹션
            p = doc.add_paragraph()
            p.add_run('고수준 컨셉: ').bold = True
            p.add_run(game_design.get('high_concept', ''))
            
            p = doc.add_paragraph()
            p.add_run('타겟 사용자층: ').bold = True
            p.add_run(game_design.get('target_audience', ''))
            
            p = doc.add_paragraph()
            p.add_run('장르: ').bold = True
            p.add_run(game_design.get('genre', ''))
            
            p = doc.add_paragraph()
            p.add_run('플랫폼: ').bold = True
            p.add_run(game_design.get('platform', ''))
            
            doc.add_paragraph()
            
            # 컨셉 세부사항
            doc.add_heading('게임 컨셉', 1)
            concept_details = game_design.get("concept_details", {})
            
            doc.add_heading('확장된 컨셉', 2)
            doc.add_paragraph(concept_details.get('extended_concept', ''))
            
            doc.add_heading('차별화 포인트', 2)
            for usp in concept_details.get("unique_selling_points", []):
                doc.add_paragraph(usp, style='List Bullet')
            
            if 'mood' in concept_details:
                doc.add_heading('분위기', 2)
                doc.add_paragraph(concept_details.get('mood', ''))
            
            # 게임플레이
            doc.add_heading('게임플레이', 1)
            gameplay = game_design.get("gameplay", {})
            
            doc.add_heading('핵심 게임플레이 루프', 2)
            doc.add_paragraph(gameplay.get('core_gameplay_loop', ''))
            
            doc.add_heading('플레이어 액션', 2)
            for action in gameplay.get("player_actions", []):
                doc.add_paragraph(action, style='List Bullet')
            
            doc.add_heading('진행 시스템', 2)
            doc.add_paragraph(gameplay.get('progression_system', ''))
            
            doc.add_heading('도전 유형', 2)
            for challenge in gameplay.get("challenge_types", []):
                doc.add_paragraph(challenge, style='List Bullet')
            
            doc.add_heading('보상 시스템', 2)
            for reward in gameplay.get("reward_systems", []):
                doc.add_paragraph(reward, style='List Bullet')
            
            doc.add_heading('게임 모드', 2)
            for mode in gameplay.get("game_modes", []):
                doc.add_paragraph(mode, style='List Bullet')
            
            if 'controls' in gameplay:
                doc.add_heading('조작 방식', 2)
                doc.add_paragraph(gameplay.get('controls', ''))
            
            doc.add_heading('독특한 메커니즘', 2)
            for mechanic in gameplay.get("unique_mechanics", []):
                doc.add_paragraph(mechanic, style='List Bullet')
            
            # 내러티브
            doc.add_heading('내러티브', 1)
            narrative = game_design.get("narrative", {})
            
            doc.add_heading('세계관', 2)
            doc.add_paragraph(narrative.get('setting', ''))
            
            doc.add_heading('배경 스토리', 2)
            doc.add_paragraph(narrative.get('background_lore', ''))
            
            doc.add_heading('주요 스토리라인', 2)
            doc.add_paragraph(narrative.get('main_plot', ''))
            
            if 'plot_structure' in narrative:
                doc.add_heading('스토리 구조', 2)
                doc.add_paragraph(narrative.get('plot_structure', ''))
            
            doc.add_heading('주요 테마', 2)
            for theme in narrative.get("themes", []):
                doc.add_paragraph(theme, style='List Bullet')
            
            # 캐릭터 섹션
            doc.add_heading('캐릭터', 2)
            for character in narrative.get("characters", []):
                doc.add_heading(character.get('name', '이름 없음'), 3)
                
                p = doc.add_paragraph()
                p.add_run('역할: ').bold = True
                p.add_run(character.get('role', ''))
                
                p = doc.add_paragraph()
                p.add_run('배경: ').bold = True
                p.add_run(character.get('background', ''))
                
                p = doc.add_paragraph()
                p.add_run('동기: ').bold = True
                p.add_run(character.get('motivation', ''))
                
                p = doc.add_paragraph()
                p.add_run('캐릭터 아크: ').bold = True
                p.add_run(character.get('arc', ''))
            
            doc.add_heading('내러티브 장치', 2)
            for device in narrative.get("narrative_devices", []):
                doc.add_paragraph(device, style='List Bullet')
            
            if 'player_agency' in narrative:
                doc.add_heading('플레이어 선택', 2)
                doc.add_paragraph(narrative.get('player_agency', ''))
            
            if 'integration_with_gameplay' in narrative:
                doc.add_heading('게임플레이와 내러티브 통합', 2)
                doc.add_paragraph(narrative.get('integration_with_gameplay', ''))
            
            # 아트 디렉션
            doc.add_heading('아트 디렉션', 1)
            art = game_design.get("art_direction", {})
            
            doc.add_heading('시각적 스타일', 2)
            doc.add_paragraph(art.get('visual_style', ''))
            
            doc.add_heading('색상 팔레트', 2)
            doc.add_paragraph(art.get('color_palette', ''))
            
            if 'art_references' in art:
                doc.add_heading('참고 아트 스타일', 2)
                for ref in art.get("art_references", []):
                    doc.add_paragraph(ref, style='List Bullet')
            
            doc.add_heading('캐릭터 디자인', 2)
            doc.add_paragraph(art.get('character_design', ''))
            
            doc.add_heading('환경 디자인', 2)
            doc.add_paragraph(art.get('environment_design', ''))
            
            doc.add_heading('UI 디자인', 2)
            doc.add_paragraph(art.get('ui_design', ''))
            
            doc.add_heading('애니메이션 스타일', 2)
            doc.add_paragraph(art.get('animation_style', ''))
            
            if 'lighting' in art:
                doc.add_heading('조명', 2)
                doc.add_paragraph(art.get('lighting', ''))
            
            doc.add_heading('사운드 디자인', 2)
            doc.add_paragraph(art.get('sound_design', ''))
            
            doc.add_heading('음악 방향', 2)
            doc.add_paragraph(art.get('music_direction', ''))
            
            # 기술 사양
            doc.add_heading('기술 사양', 1)
            tech = game_design.get("technical_specs", {})
            
            doc.add_heading('타겟 플랫폼', 2)
            for platform in tech.get("target_platforms", []):
                doc.add_paragraph(platform, style='List Bullet')
            
            if 'minimum_specs' in tech:
                doc.add_heading('최소 사양', 2)
                doc.add_paragraph(tech.get('minimum_specs', ''))
            
            if 'recommended_specs' in tech:
                doc.add_heading('권장 사양', 2)
                doc.add_paragraph(tech.get('recommended_specs', ''))
            
            doc.add_heading('게임 엔진', 2)
            doc.add_paragraph(tech.get('engine', ''))
            
            doc.add_heading('핵심 기술', 2)
            for technology in tech.get("key_technologies", []):
                doc.add_paragraph(technology, style='List Bullet')
            
            if 'networking' in tech:
                doc.add_heading('네트워킹', 2)
                doc.add_paragraph(tech.get('networking', ''))
            
            if 'performance_targets' in tech:
                doc.add_heading('성능 목표', 2)
                doc.add_paragraph(tech.get('performance_targets', ''))
            
            doc.add_heading('개발 도전 과제', 2)
            for challenge in tech.get("development_challenges", []):
                doc.add_paragraph(challenge, style='List Bullet')
            
            doc.add_heading('에셋 요구사항', 2)
            doc.add_paragraph(tech.get('asset_requirements', ''))
            
            # 수익화 계획
            doc.add_heading('수익화 계획', 1)
            monetization = game_design.get("monetization", {})
            
            doc.add_heading('수익화 모델', 2)
            doc.add_paragraph(monetization.get('monetization_model', ''))
            
            doc.add_heading('가격 책정', 2)
            doc.add_paragraph(monetization.get('price_point', ''))
            
            doc.add_heading('인앱 구매', 2)
            doc.add_paragraph(monetization.get('in_app_purchases', ''))
            
            doc.add_heading('DLC/확장팩', 2)
            doc.add_paragraph(monetization.get('dlc_expansion', ''))
            
            doc.add_heading('플레이어 유지 전략', 2)
            for strategy in monetization.get("player_retention_strategies", []):
                doc.add_paragraph(strategy, style='List Bullet')
            
            doc.add_heading('시장 분석', 2)
            doc.add_paragraph(monetization.get('market_analysis', ''))
            
            # 개발 로드맵
            doc.add_heading('개발 로드맵', 1)
            roadmap = game_design.get("development_roadmap", {})
            
            # 스토리라인 섹션 (있는 경우)
            if "storyline" in game_design:
                storyline = game_design.get("storyline", {})
                
                doc.add_heading('스토리라인', 1)
                
                doc.add_heading('스토리 제목', 2)
                doc.add_paragraph(storyline.get('title', ''))
                
                doc.add_heading('전제', 2)
                doc.add_paragraph(storyline.get('premise', ''))
                
                doc.add_heading('스토리 아크', 2)
                doc.add_paragraph(storyline.get('story_arc', ''))
                
                # 챕터 개요
                doc.add_heading('챕터 개요', 2)
                for chapter in storyline.get('chapter_outlines', []):
                    doc.add_heading(f"챕터 {chapter.get('chapter_number')}: {chapter.get('title')}", 3)
                    doc.add_paragraph(chapter.get('synopsis', ''))
                    
                    p = doc.add_paragraph()
                    p.add_run("목표:").bold = True
                    for goal in chapter.get('goals', []):
                        doc.add_paragraph(goal, style='List Bullet')
                    
                    p = doc.add_paragraph()
                    p.add_run("주요 위치:").bold = True
                    for location in chapter.get('key_locations', []):
                        doc.add_paragraph(location, style='List Bullet')
                    
                    p = doc.add_paragraph()
                    p.add_run("주요 사건:").bold = True
                    for event in chapter.get('key_events', []):
                        doc.add_paragraph(event, style='List Bullet')
                    
                    p = doc.add_paragraph()
                    p.add_run("도전 과제:").bold = True
                    for challenge in chapter.get('challenges', []):
                        doc.add_paragraph(challenge, style='List Bullet')
                
                # 챕터 상세 내용 섹션 추가
                doc.add_heading('챕터 상세 내용', 2)
                for chapter in storyline.get('chapter_details', []):
                    doc.add_heading(f"챕터 {chapter.get('chapter_number')}: {chapter.get('title')}", 3)
                    
                    # 상세 시놉시스
                    doc.add_paragraph(chapter.get('detailed_synopsis', ''))
                    
                    # 오프닝 씬
                    p = doc.add_paragraph()
                    p.add_run('오프닝 씬: ').bold = True
                    doc.add_paragraph(chapter.get('opening_scene', ''))
                    
                    # 주요 사건
                    p = doc.add_paragraph()
                    p.add_run('주요 사건:').bold = True
                    for event in chapter.get('key_events', []):
                        p = doc.add_paragraph()
                        p.add_run(f"{event.get('event_title', '')}: ").bold = True
                        p.add_run(event.get('description', ''))
                    
                    # 클라이맥스
                    p = doc.add_paragraph()
                    p.add_run('클라이맥스: ').bold = True
                    doc.add_paragraph(chapter.get('climax', ''))
                    
                    # 엔딩
                    p = doc.add_paragraph()
                    p.add_run('엔딩: ').bold = True
                    doc.add_paragraph(chapter.get('ending', ''))
            
            # 문서 저장
            doc.save(file_path)
            print(f"Word 문서 생성 완료: {file_path}")
            
        except Exception as e:
            print(f"Word 문서 생성 중 오류 발생: {str(e)}")
            raise